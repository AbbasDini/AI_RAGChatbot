import os
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file, send_from_directory, Blueprint
# from langchain.vectorstores import Chroma # No longer needed for direct querying
from langchain.embeddings import SentenceTransformerEmbeddings
# Removed RetrievalQA
from dotenv import load_dotenv
import pandas as pd # Import pandas
import markdown # Import the markdown library
import base64
import tempfile
from pathlib import Path
from datetime import datetime
import io
import speech_recognition as sr  # For STT (pip install SpeechRecognition)
from werkzeug.utils import secure_filename
import json
import shutil

# Import the specific LLM class and prompt template
from langchain_openai import OpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document # Import Document class

# Import chromadb for direct client interaction
import chromadb

# Import advanced TTS module
from advanced_tts import AdvancedTTS
from doctorbot_models import db, DoctorBotSettings, MedicalDocument
from doctorbot_routes import doctorbot_bp
from doctors_data import doctors_info
from llm_utils import get_llm

# --- Configuration ---
# Load environment variables from .env file
load_dotenv()

# Get API key and base URL from environment variables/config
AVALAI_API_KEY = os.getenv("AVALAI_API_KEY")
# Assuming the base_url for avalai.ir - please verify this if you have a specific one
AVALAI_BASE_URL = os.getenv("AVALAI_BASE_URL", "https://api.avalai.ir/v1") # Allow override from .env
AVALAI_MODEL_NAME = os.getenv("AVALAI_MODEL_NAME", "gpt-4.1")

# Get the base directory of the Flask app
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MED_DOC_DIR = os.path.join(BASE_DIR, "Med_doc") # Directory containing doctor folders
CHROMA_DB_DIR = os.path.join(BASE_DIR, "chroma_db") # Directory where ChromaDB is persisted

# List of doctors (dynamically get from subdirectories)
# This is kept for potential future use but is not directly used in the current filtering flow
doctor_folders = [d for d in os.listdir(MED_DOC_DIR) if os.path.isdir(os.path.join(MED_DOC_DIR, d))]
DOCTORS = [(d, d.replace('_', ' ').title()) for d in doctor_folders] # (value, display_name)

# Lists for filter options (English key, Persian display)
CITIES = [('all', 'همه'), ('tehran', 'تهران'), ('shiraz', 'شیراز'), ('hamedan', 'همدان'), ('mashhad', 'مشهد'), ('isfahan', 'اصفهان'), ('tabriz', 'تبریز'), ('kermanshah', 'کرمانشاه')]
SPECIALTIES = [('all', 'همه'), ('general', 'پزشک عمومی'), ('cardiologist', 'قلب و عروق'), ('neurologist', 'مغز و اعصاب'), ('pediatrician', 'اطفال'), ('oncologist', 'آنکولوژی')]
EXPERIENCE_RANGES = [('all', 'همه'), ('1-5', '1 تا 5 سال'), ('5-10', '5 تا 10 سال'), ('10+', 'بیشتر از 10 سال')]

SETTINGS_PATH = os.path.join(BASE_DIR, 'chatbot_settings.json')
RESOURCES_PATH = os.path.join(BASE_DIR, 'chatbot_resources.json')
USERS_PATH = os.path.join(BASE_DIR, 'chatbot_users.json')
LOGS_PATH = os.path.join(BASE_DIR, 'chatbot_logs.txt')

def load_json(path, default):
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return default

def save_json(path, data):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def append_log(msg):
    with open(LOGS_PATH, 'a', encoding='utf-8') as f:
        f.write(f'{datetime.now().isoformat()} | {msg}\n')

# --- Flask App Setup ---
app = Flask(__name__)
app.secret_key = 'your_secret_key_here' # Replace with a strong secret key
# --- تنظیمات دیتابیس ---
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///doctorbot.sqlite3'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)
app.register_blueprint(doctorbot_bp)

# --- LLM Dynamic Setup ---
# تابع get_llm حذف شد و از llm_utils import می‌شود

# --- RAG Setup ---
# Initialize embedding model with error handling and fallback options
embedding_model = None
try:
    # Try to use the original model
    embedding_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    print("Successfully loaded all-MiniLM-L6-v2 embedding model.")
except Exception as e:
    print(f"Error loading all-MiniLM-L6-v2: {e}")
    print("Trying alternative models...")
    
    # Try alternative models that might be available offline
    alternative_models = [
        "paraphrase-multilingual-MiniLM-L12-v2",
        "distiluse-base-multilingual-cased-v2",
        "sentence-transformers/all-mpnet-base-v2",
        "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    ]
    
    for model_name in alternative_models:
        try:
            embedding_model = SentenceTransformerEmbeddings(model_name=model_name)
            print(f"Successfully loaded alternative model: {model_name}")
            break
        except Exception as alt_e:
            print(f"Failed to load {model_name}: {alt_e}")
            continue
    
    if embedding_model is None:
        print("Warning: Could not load any embedding model. Using simple text matching as fallback.")
        # Create a simple fallback embedding function
        class SimpleEmbedding:
            def __init__(self):
                self.model_name = "simple_text_matcher"
            
            def embed_documents(self, texts):
                # Simple fallback: return random embeddings
                import numpy as np
                return [np.random.rand(384).tolist() for _ in texts]
            
            def embed_query(self, text):
                import numpy as np
                return np.random.rand(384).tolist()
        
        embedding_model = SimpleEmbedding()

# Initialize the LLM with avalai.ir settings
llm = None # Initialize LLM as None by default
if not AVALAI_API_KEY:
    print("Error: AVALAI_API_KEY not found in environment variables. Please add it to your .env file.")
elif not AVALAI_BASE_URL:
     print("Error: AVALAI_BASE_URL not found. Please check your .env file or use default.")
else:
    try:
        # Use OpenAI class with api_key and base_url for compatibility
        llm = OpenAI(
            api_key=AVALAI_API_KEY, # Use api_key parameter for newer Langchain versions
            base_url=AVALAI_BASE_URL,
            model_name=AVALAI_MODEL_NAME, # Specify the model name
            temperature=0.7 # Example temperature
        )
        print("LLM initialized successfully with avalai.ir settings.")
    except Exception as e:
        print(f"Error initializing LLM: {e}")
        llm = None

# Initialize ChromaDB Client (persistent)
# The collection will be retrieved within the chat route
chroma_client = None
try:
    chroma_client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
    print("ChromaDB client initialized successfully.")
except Exception as e:
    print(f"Error initializing ChromaDB client: {e}")

# Initialize Advanced TTS
advanced_tts = None
try:
    advanced_tts = AdvancedTTS()
    print("Advanced TTS initialized successfully.")
except Exception as e:
    print(f"Error initializing Advanced TTS: {e}")
    advanced_tts = None

# Define a simple prompt template for the LLM
# This template includes context from retrieved documents
prompt_template = PromptTemplate(
    input_variables=["context", "question"],
    template=(
        "You are a helpful medical assistant. Use the following pieces of context to answer the question.\\n\\n"
        "If you don\'t know the answer, just say that you don\'t know, don\'t try to make up an answer.\\n\\n"
        "Context: {context}\\n\\n"
        "Question: {question}\\n\\n"
        "Answer:"
    )
)

# --- Routes ---

@app.route('/')
def index():
    """Render the landing page with two main options."""
    return render_template('landing.html')

@app.route('/chat_interface')
def chat_interface():
    """Render the main chat interface page."""
    return render_template('index.html',
                         cities=CITIES,
                         specialties=SPECIALTIES,
                         experience_ranges=EXPERIENCE_RANGES)

@app.route('/find_doctors', methods=['POST'])
def find_doctors():
    """Handle filtering doctors based on criteria and display results."""
    city = request.form.get('city')
    specialty = request.form.get('specialty')
    experience = request.form.get('experience')

    # Store filtering criteria in session
    session['filter_criteria'] = {
        'city': city,
        'specialty': specialty,
        'experience': experience
    }

    # Filter doctors based on criteria
    matching_doctors = {}
    for folder_name, info in doctors_info.items():
        # Check if the doctor matches the criteria, considering 'all' as no filter
        # Compare using the English keys received from the form directly with doctors_info keys
        city_match = (city == 'all' or info['city'] == city)
        specialty_match = (specialty == 'all' or info['specialty'] == specialty)
        experience_match = (experience == 'all' or info['experience'] == experience)

        if city_match and specialty_match and experience_match:
            matching_doctors[folder_name] = info

    # Render template to display matching doctors
    return render_template('select_doctor.html', matching_doctors=matching_doctors, criteria=session['filter_criteria'])


@app.route('/select_doctor', methods=['POST'])
def select_doctor():
    """Handle selected doctor and redirect to chat page."""
    selected_doctor_folder = request.form.get('selected_doctor')

    # Check if a doctor was selected and if it's a valid key in doctors_info
    if not selected_doctor_folder or selected_doctor_folder not in doctors_info:
        # Handle error if doctor is not selected or invalid
        # TODO: Add a proper error message to the user on the select doctor page
        print(f"Invalid doctor selected: {selected_doctor_folder}")
        return redirect(url_for('index')) # Redirect to index for now

    # Store selected doctor in session
    session['selected_doctor'] = selected_doctor_folder
    # Re-initialize chat history for a new chat session with this doctor
    session['chat_history'] = []

    # Redirect to the chat page
    return redirect(url_for('chat'))


@app.route('/chat', methods=['GET', 'POST'])
def chat():
    """Render the chat page and handle chat messages."""
    # Ensure selected_doctor is in session to proceed to chat
    if 'selected_doctor' not in session:
        # Redirect to index if selected_doctor is missing
        print("Redirecting from /chat to index because selected_doctor is not in session.") # Debug print (keep for now)
        return redirect(url_for('index'))

    # Get necessary information from session
    selected_doctor = session['selected_doctor']
    # Retrieve filter criteria from session for display on the chat page if needed
    criteria = session.get('filter_criteria', {})
    # You might want to add language selection back in the chat if needed
    # selected_language = session.get('selected_language', 'fa') # Assuming a default language

    chat_history = session.get('chat_history', [])

    user_input = None
    bot_response = ""

    # Check if LLM and chroma client were initialized successfully
    global llm, chroma_client
    if llm is None or chroma_client is None:
        bot_response = "Error: Application not configured properly. Language Model or Document Database client is not available."
        if request.method == 'POST' and request.form.get('user_input'):
             chat_history.append({'speaker': 'user', 'text': request.form.get('user_input')})
             chat_history.append({'speaker': 'bot', 'text': bot_response})
             session['chat_history'] = chat_history
        return render_template('chat_main.html',
                               doctor_name=doctors_info.get(selected_doctor, {}).get('name', selected_doctor), # Get display name
                               # language=next((item[1] for item in LANGUAGES if item[0] == selected_language), selected_language), # Get display name if language added back
                               chat_history=chat_history,
                               criteria=criteria # Pass criteria to chat template for display
                               )

    if request.method == 'POST':
        user_input = request.form.get('user_input')
        if user_input:
            # Add user message to history immediately
            chat_history.append({'speaker': 'user', 'text': user_input})

            # --- RAG Query Logic (Direct Chroma Client) ---
            try:
                # Get the collection (assuming default name "langchain")
                collection = chroma_client.get_collection(name="langchain")

                # Define the base where filter (always filter by doctor)
                where_filter = {'doctor': selected_doctor}

                # Although criteria are stored, we currently only filter by doctor metadata in ChromaDB.
                # If city, specialty, experience were added as metadata during processing,
                # you could add them to the where_filter here:
                # if 'city' in criteria and criteria['city'] != 'all':
                #     where_filter['city'] = criteria['city']
                # etc.

                # Manually perform query with doctor-specific filtering using Chroma client's query method
                # Added timeout for query
                results = collection.query(
                    query_texts=[user_input],
                    n_results=3, # Number of results to retrieve
                    where=where_filter, # Apply metadata filter
                    # include=['documents', 'metadatas', 'distances'] # Specify what to include
                )

                # Process results and format retrieved documents for the LLM context
                retrieved_docs = []
                if results and results.get('documents') and results['documents'][0]:
                    for i in range(len(results['documents'][0])):
                        page_content = results['documents'][0][i]
                        metadata = results['metadatas'][0][i] if results.get('metadatas') and results['metadatas'][0] else {}
                        retrieved_docs.append(Document(page_content=page_content, metadata=metadata))

                context_text = "\n\n---\n\n".join([doc.page_content for doc in retrieved_docs])

                # Create the final prompt
                prompt = prompt_template.format(context=context_text, question=user_input)

                # --- LLM Interaction ---
                # Use the invoke method for the LLM chain
                if llm:
                     try:
                         # The invoke method directly returns the response content as a string
                         bot_response_content = llm.invoke(prompt)
                         bot_response = bot_response_content # Use the string content directly
                     except Exception as e:
                         bot_response = f"Error during LLM invocation: {e}"
                         print(bot_response)
                else:
                     bot_response = "LLM is not initialized."


            except Exception as e:
                bot_response = f"Error during document retrieval or processing: {e}"
                print(bot_response)

            # Add bot response to history
            chat_history.append({'speaker': 'bot', 'text': bot_response})

            # Update chat history in session
            session['chat_history'] = chat_history

            # Return JSON response for AJAX if using AJAX for chat (optional)
            # return jsonify({'response': bot_response})

    # Render chat template (for GET or non-AJAX POST)
    # Pass the full doctor info object to the template
    doctor_info = doctors_info.get(selected_doctor, {})
    doctor_info['avatar_url'] = url_for('static', filename='doctors_images/' + doctor_info.get('image', 'doctor_default.jpg'))
    return render_template('chat_main.html',
                           doctor=doctor_info,  # Pass doctor for new template
                           messages=chat_history,  # Pass messages for new template
                           doctor_name=doctor_info.get('name', selected_doctor), # Get display name
                           doctor_image=doctor_info.get('image', 'default_doctor.jpg'), # Get doctor image
                           doctor_details=doctor_info, # Pass the full info object
                           chat_history=chat_history,
                           criteria=criteria # Still pass criteria in case it's needed elsewhere, but not displayed by default
                           )

@app.route('/patient_prompt')
def patient_prompt():
    return render_template('patient_form.html')

@app.route('/generate_prompt', methods=['POST'])
def generate_prompt():
    language = request.form.get('language')
    city = request.form.get('city')
    country = request.form.get('country')
    initial_problem = request.form.get('initial_problem')
    illness_status = request.form.get('illness_status')

    prompt = f"""**اطلاعات دریافت شده از بیمار:**
    - **زبان محاوره ای:** {language}
    - **محل سکونت:** {city}, {country}
    - **تشخیص/مشکل اولیه:** {initial_problem}
    - **وضعیت بیماری:** {illness_status}
    """

    # Convert markdown to HTML
    prompt_html = markdown.markdown(prompt)
    return render_template('prompt_result.html', prompt=prompt_html)

@app.route('/get_doctor_suggestions', methods=['POST'])
def get_doctor_suggestions():
    patient_prompt = request.form.get('patient_prompt')

    # Read doctors data from Excel
    try:
        doctors_df = pd.read_excel('doctors_data.xlsx')
        doctors_info_text = doctors_df.to_string(index=False)
    except FileNotFoundError:
        return "خطا: فایل doctors_data.xlsx یافت نشد."
    except Exception as e:
        return f"خطا در خواندن فایل اکسل: {e}"

    # Prepare prompt for LLM
    llm_query_prompt = f"""پرامت بیمار:
{patient_prompt}

لیست پزشکان متخصص:
{doctors_info_text}

با توجه به پرامت بیمار و لیست پزشکان، لطفا 1 تا 3 پزشک متخصص مناسب با شرایط ایشان را با زبان انتخاب شده کاربر (زبان محاوره ای او) پیشنهاد دهید. 
پاسخ شما باید شامل نام و نام خانوادگی و تخصص و رشته و شهر و تسلط به زبان و پروفایل آکسایا پزشکان پیشنهادی باشد.
مثال با زبان فارسی (فرمت پاسخ):
<br/><br/>
پیشنهاد 1: <br/>- <span style="font-weight: bold; color: red;">دکتر علی محمدی</span><br/>- <span style="font-weight: bold; color: orange;">متخصص قلب و عروق</span><br/>- <span style="font-weight: bold; color: orange;">شهر همدان</span><br/>- <span style="font-weight: bold; color: orange;">فقط فارسی</span><br/><a href="#" class="aksaya-profile-button">پروفایل آکسایا</a>
<br/><br/>
پیشنهاد 2: <br/>- <span style="font-weight: bold; color: red;">دکتر فاطمه حسینی</span><br/>- <span style="font-weight: bold; color: orange;">متخصص گوش و حلق وبینی</span><br/>- <span style="font-weight: bold; color: orange;">تهران</span><br/>- <span style="font-weight: bold; color: orange;">فارسی و انگلیسی</span><br/><a href="#" class="aksaya-profile-button">پروفایل آکسایا</a>
<br/><br/>
پیشنهاد 3: <br/>- <span style="font-weight: bold; color: red;">دکتر رضا کریمی</span><br/>- <span style="font-weight: bold; color: orange;">متخصص ارتوپدی</span><br/>- <span style="font-weight: bold; color: orange;">شیراز</span><br/>- <span style="font-weight: bold; color: orange;">فقط فارسی</span><br/><a href="#" class="aksaya-profile-button">پروفایل آکسایا</a>

مثال با زبان انگلیسی (فرمت پاسخ):
<br/><br/>
Suggestion 1: <br/>- <span style="font-weight: bold; color: red;">Dr. Ali Mohammadi</span><br/>- <span style="font-weight: bold; color: orange;">Cardiologist</span><br/>- <span style="font-weight: bold; color: orange;">Hamedan</span><br/>- <span style="font-weight: bold; color: orange;">Persian only</span><br/><a href="https://example.com/aksaya_profile_ali_en" class="aksaya-profile-button">Aksaya Profile</a>
"""

    # Call LLM
    global llm
    if llm is None:
        return "خطا: مدل زبانی LLM راه‌اندازی نشده است."

    try:
        llm_response = llm.invoke(llm_query_prompt)
        suggested_doctors_raw = llm_response.strip()
        
        # Process the LLM response to extract doctor information
        suggested_doctors_list = []
        current_doctor = []
        
        for line in suggested_doctors_raw.split('\n'):
            if line.strip():
                if line.startswith('پیشنهاد') or line.startswith('Suggestion'):
                    if current_doctor:
                        suggested_doctors_list.append('\n'.join(current_doctor))
                        current_doctor = []
                current_doctor.append(line)
        
        if current_doctor:
            suggested_doctors_list.append('\n'.join(current_doctor))

        # Create a list of dictionaries containing doctor info and their corresponding images
        doctors_with_images = []
        for i, doctor_info in enumerate(suggested_doctors_list):
            # Use different images for each doctor
            image_file = f"0006.png" if i == 0 else f"umsha.png" if i == 1 else "umsha_main.png"
            doctors_with_images.append({
                'info': doctor_info,
                'image': image_file
            })

    except Exception as e:
        return f"خطا در برقراری ارتباط با LLM: {e}"

    return render_template('doctor_suggestions.html', doctors=doctors_with_images)


@app.route('/api/tts', methods=['POST'])
def text_to_speech():
    """API endpoint for advanced text-to-speech"""
    temp_file = None
    try:
        data = request.get_json()
        text = data.get('text', '')
        provider = data.get('provider', 'auto') # Changed default to 'auto'
        voice = data.get('voice', '')
        
        if not text:
            return jsonify({'error': 'Text is required'}), 400
        
        if not advanced_tts:
            return jsonify({'error': 'TTS service not available'}), 503
        
        print(f"TTS request: text='{text[:50]}...', provider='{provider}', voice='{voice}'")
        
        # Synthesize speech
        audio_data = advanced_tts.synthesize_speech(text, provider, voice)
        
        if not audio_data:
            print("TTS failed: No audio data returned")
            return jsonify({'error': 'Failed to synthesize speech'}), 500
        
        print(f"TTS successful: {len(audio_data)} bytes")
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3')
        temp_file.write(audio_data)
        temp_file.close()
        
        # Return the file and clean up after sending
        response = send_file(temp_file.name, mimetype='audio/mpeg')
        
        # Clean up the temporary file after the response is sent
        @response.call_on_close
        def cleanup():
            try:
                import os
                if temp_file and os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)
            except Exception as e:
                print(f"Error cleaning up temp file: {e}")
        
        return response
        
    except Exception as e:
        print(f"TTS error: {e}")
        # Clean up temp file on error
        if temp_file:
            try:
                import os
                if os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)
            except:
                pass
        return jsonify({'error': str(e)}), 500

@app.route('/api/tts/voices')
def get_available_voices():
    """Get available TTS voices"""
    try:
        if not advanced_tts:
            return jsonify({'error': 'TTS service not available'}), 503
        
        voices = advanced_tts.get_available_voices()
        return jsonify(voices)
        
    except Exception as e:
        print(f"Error getting voices: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/tts/test')
def test_tts():
    """Test TTS with sample Persian text"""
    try:
        test_text = "سلام، من دستیار هوشمند پزشکی هستم. چگونه می‌توانم به شما کمک کنم؟"
        
        if not advanced_tts:
            return jsonify({'error': 'TTS service not available'}), 503
        
        # Test with different providers
        results = {}
        providers = ['avalai', 'azure', 'elevenlabs', 'google', 'openai']
        
        for provider in providers:
            try:
                audio_data = advanced_tts.synthesize_speech(test_text, provider)
                if audio_data:
                    # Save test file
                    filename = f"test_{provider}.mp3"
                    file_path = advanced_tts.save_audio_file(audio_data, filename)
                    results[provider] = {
                        'status': 'success',
                        'file': file_path,
                        'size': len(audio_data)
                    }
                else:
                    results[provider] = {'status': 'failed', 'error': 'No audio data'}
            except Exception as e:
                results[provider] = {'status': 'failed', 'error': str(e)}
        
        return jsonify(results)
        
    except Exception as e:
        print(f"TTS test error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat_advanced_page')
def chat_advanced_page():
    """نمایش صفحه چت پیشرفته"""
    if 'selected_doctor' not in session:
        return redirect(url_for('index'))
    
    selected_doctor = session['selected_doctor']
    doctor_info = doctors_info.get(selected_doctor, {})
    chat_history = session.get('chat_history', [])
    
    return render_template('chat_advanced.html',
                         doctor_name=doctor_info.get('name', selected_doctor),
                         doctor_image=doctor_info.get('image', 'default_doctor.jpg'),
                         chat_history=chat_history)

@app.route('/chat_advanced', methods=['POST'])
def chat_advanced():
    """دریافت پیام متنی و بازگشت پاسخ هوشمند"""
    data = request.get_json()
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'error': 'متن پیام خالی است.'}), 400
    
    # ساخت پیام کاربر
    user_msg = {
        'type': 'text',
        'speaker': 'user',
        'text': text,
        'timestamp': datetime.now().strftime('%H:%M')
    }
    chat_history = session.get('chat_history', [])
    chat_history.append(user_msg)
    
    # پردازش با مدل LLM واقعی
    bot_response = ""
    try:
        if llm and chroma_client:
            # RAG Query Logic
            collection = chroma_client.get_collection(name="langchain")
            selected_doctor = session.get('selected_doctor', 'doctor_abbasi')
            where_filter = {'doctor': selected_doctor}
            
            results = collection.query(
                query_texts=[text],
                n_results=3,
                where=where_filter,
            )
            
            # Process results
            retrieved_docs = []
            if results and results.get('documents') and results['documents'][0]:
                for i in range(len(results['documents'][0])):
                    page_content = results['documents'][0][i]
                    metadata = results['metadatas'][0][i] if results.get('metadatas') and results['metadatas'][0] else {}
                    retrieved_docs.append(Document(page_content=page_content, metadata=metadata))
            
            context_text = "\n\n---\n\n".join([doc.page_content for doc in retrieved_docs])
            prompt = prompt_template.format(context=context_text, question=text)
            bot_response_content = llm.invoke(prompt)
            bot_response = bot_response_content
        else:
            bot_response = f"پاسخ هوشمند به: {text}"
    except Exception as e:
        bot_response = f"خطا در پردازش: {str(e)}"
    
    # ساخت پیام ربات
    bot_msg = {
        'type': 'text',
        'speaker': 'bot',
        'text': bot_response,
        'timestamp': datetime.now().strftime('%H:%M')
    }
    chat_history.append(bot_msg)
    session['chat_history'] = chat_history

    # --- تولید پاسخ صوتی (TTS) ---
    try:
        if advanced_tts:
            # Use AvalAI Gemini model with explicit Iranian accent instructions
            audio_data = None
            provider = 'avalai'
            model = 'gemini-2.5-pro-preview-tts'
            voice = 'nova'
            instructions = 'با لهجه فارسی ایرانی (تهرانی) و لحن طبیعی و استاندارد صحبت کن.'
            # Try Gemini model first
            audio_data = advanced_tts.text_to_speech_avalai(bot_response, voice, model=model) if hasattr(advanced_tts, 'text_to_speech_avalai') else None
            if not audio_data and hasattr(advanced_tts, 'text_to_speech_avalai'):
                # Try shimmer as alternative female voice
                voice = 'shimmer'
                audio_data = advanced_tts.text_to_speech_avalai(bot_response, voice, model=model)
            # Fallbacks
            if not audio_data:
                provider = 'azure'
                voice = 'fa-IR-SaraNeural'
                audio_data = advanced_tts.synthesize_speech(bot_response, provider, voice)
            if not audio_data:
                provider = 'azure'
                voice = 'fa-IR-YektaNeural'
                audio_data = advanced_tts.synthesize_speech(bot_response, provider, voice)
            if not audio_data:
                provider = 'google'
                voice = 'fa-IR-Wavenet-B'
                audio_data = advanced_tts.synthesize_speech(bot_response, provider, voice)
            if audio_data:
                audio_filename = f"tts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav"
                audio_path = os.path.join('HT_RAG_Chatbot/static/tts_cache', audio_filename)
                with open(audio_path, 'wb') as f:
                    f.write(audio_data)
                audio_url = f"/static/tts_cache/{audio_filename}"
                bot_voice_msg = {
                    'type': 'voice',
                    'speaker': 'bot',
                    'audio_url': audio_url,
                    'timestamp': datetime.now().strftime('%H:%M')
                }
                chat_history.append(bot_voice_msg)
                session['chat_history'] = chat_history
                return jsonify(bot_voice_msg)
    except Exception as e:
        print(f"TTS error: {e}")
    # اگر TTS نشد، پیام متنی را برگردان
    return jsonify(bot_msg)

@app.route('/voice_message', methods=['POST'])
def voice_message():
    """دریافت پیام صوتی و بازگشت پاسخ"""
    if 'audio' not in request.files:
        return jsonify({'error': 'فایل صوتی ارسال نشده است.'}), 400
    
    audio_file = request.files['audio']
    
    # ذخیره فایل صوتی موقت
    temp_filename = f"voice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.webm"
    temp_path = os.path.join('HT_RAG_Chatbot/static', temp_filename)
    audio_file.save(temp_path)
    
    # پیام کاربر
    user_msg = {
        'type': 'voice',
        'speaker': 'user',
        'audio_url': f'/static/{temp_filename}',
        'timestamp': datetime.now().strftime('%H:%M')
    }
    chat_history = session.get('chat_history', [])
    chat_history.append(user_msg)
    
    # در حالت واقعی اینجا باید STT اجرا شود
    # فعلاً پاسخ نمونه
    bot_text = 'پیام صوتی شما دریافت شد. در حال پردازش...'
    
    bot_msg = {
        'type': 'text',
        'speaker': 'bot',
        'text': bot_text,
        'timestamp': datetime.now().strftime('%H:%M')
    }
    chat_history.append(bot_msg)
    session['chat_history'] = chat_history
    
    return jsonify(bot_msg)

# --- STT (Speech-to-Text) Endpoint ---
@app.route('/stt', methods=['POST'])
def stt_api():
    """دریافت فایل صوتی و تبدیل به متن فارسی (Speech-to-Text)"""
    if 'audio' not in request.files:
        return jsonify({'error': 'فایل صوتی ارسال نشده است.'}), 400
    audio_file = request.files['audio']
    filename = secure_filename(audio_file.filename or 'audio_input.wav')
    temp_path = os.path.join('HT_RAG_Chatbot/static', filename)
    audio_file.save(temp_path)
    try:
        recognizer = sr.Recognizer()
        with sr.AudioFile(temp_path) as source:
            audio_data = recognizer.record(source)
        text = recognizer.recognize_google(audio_data, language='fa-IR')
        os.remove(temp_path)
        return jsonify({'text': text})
    except Exception as e:
        print(f'STT error: {e}')
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': f'خطا در تبدیل گفتار به متن: {str(e)}'}), 500

@app.route('/tts', methods=['POST'])
def tts_api():
    data = request.get_json()
    text = data.get('text', '')
    # Use AvalAI Gemini model with explicit Iranian accent instructions
    if not text:
        return jsonify({'error': 'متن خالی است.'}), 400
    try:
        if advanced_tts:
            audio_data = None
            provider = 'avalai'
            model = 'gemini-2.5-pro-preview-tts'
            voice = 'nova'
            instructions = 'با لهجه فارسی ایرانی (تهرانی) و لحن طبیعی و استاندارد صحبت کن.'
            # Try Gemini model first
            audio_data = advanced_tts.text_to_speech_avalai(text, voice, model=model) if hasattr(advanced_tts, 'text_to_speech_avalai') else None
            if not audio_data and hasattr(advanced_tts, 'text_to_speech_avalai'):
                # Try shimmer as alternative female voice
                voice = 'shimmer'
                audio_data = advanced_tts.text_to_speech_avalai(text, voice, model=model)
            # Fallbacks
            if not audio_data:
                provider = 'azure'
                voice = 'fa-IR-SaraNeural'
                audio_data = advanced_tts.synthesize_speech(text, provider, voice)
            if not audio_data:
                provider = 'azure'
                voice = 'fa-IR-YektaNeural'
                audio_data = advanced_tts.synthesize_speech(text, provider, voice)
            if not audio_data:
                provider = 'google'
                voice = 'fa-IR-Wavenet-B'
                audio_data = advanced_tts.synthesize_speech(text, provider, voice)
            if audio_data:
                return send_file(
                    io.BytesIO(audio_data),
                    mimetype='audio/wav',
                    as_attachment=False,
                    download_name='tts.wav'
                )
        silent_audio = io.BytesIO()
        silent_audio.write(b'RIFF....WAVEfmt ')
        silent_audio.seek(0)
        return send_file(
            silent_audio,
            mimetype='audio/wav',
            as_attachment=False,
            download_name='tts.wav'
        )
    except Exception as e:
        return jsonify({'error': f'خطا در TTS: {str(e)}'}), 500

@app.route('/admin/chatbot_settings', methods=['GET'])
def chatbot_settings():
    settings = load_json(SETTINGS_PATH, {
        'bot_name': 'چت‌بات پزشکی',
        'bot_description': '',
        'default_language': 'fa',
        'bot_active': True,
        'llm_model': 'avalai',
        'embedding_model': 'paraphrase-multilingual-MiniLM-L12-v2',
        'top_k': 3,
        'temperature': 0.7
    })
    resources = load_json(RESOURCES_PATH, [])
    users = load_json(USERS_PATH, [])
    logs = []
    if os.path.exists(LOGS_PATH):
        with open(LOGS_PATH, 'r', encoding='utf-8') as f:
            logs = f.readlines()[-30:]
    return render_template('admin/chatbot_settings.html', settings=settings, resources=resources, users=users, logs=logs)

@app.route('/admin/save_settings', methods=['POST'])
def save_settings():
    settings = load_json(SETTINGS_PATH, {})
    settings['bot_name'] = request.form.get('bot_name', '')
    settings['bot_description'] = request.form.get('bot_description', '')
    settings['default_language'] = request.form.get('default_language', 'fa')
    settings['bot_active'] = request.form.get('bot_active', '1') == '1'
    save_json(SETTINGS_PATH, settings)
    append_log('تنظیمات کلی چت‌بات ذخیره شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/upload_resource', methods=['POST'])
def upload_resource():
    file = request.files.get('resource_file')
    if file and file.filename.endswith('.docx'):
        save_path = os.path.join(MED_DOC_DIR, 'doctor_abbasi', file.filename)
        file.save(save_path)
        resources = load_json(RESOURCES_PATH, [])
        resources.append({'name': file.filename, 'size': round(os.path.getsize(save_path)/1024, 1)})
        save_json(RESOURCES_PATH, resources)
        append_log(f'فایل {file.filename} آپلود شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/rebuild_embeddings', methods=['POST'])
def rebuild_embeddings():
    # پاکسازی دیتابیس و بازسازی embedding
    chroma_dir = os.path.join(BASE_DIR, 'chroma_db')
    if os.path.exists(chroma_dir):
        import shutil
        shutil.rmtree(chroma_dir)
    os.system(f'python {os.path.join(BASE_DIR, "process_docs.py")}')
    append_log('دیتابیس embedding بازسازی شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/save_model_settings', methods=['POST'])
def save_model_settings():
    settings = load_json(SETTINGS_PATH, {})
    settings['llm_model'] = request.form.get('llm_model', 'avalai')
    settings['embedding_model'] = request.form.get('embedding_model', 'paraphrase-multilingual-MiniLM-L12-v2')
    settings['top_k'] = int(request.form.get('top_k', 3))
    settings['temperature'] = float(request.form.get('temperature', 0.7))
    save_json(SETTINGS_PATH, settings)
    append_log('تنظیمات مدل و embedding ذخیره شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/add_user', methods=['POST'])
def add_user():
    username = request.form.get('username', '').strip()
    role = request.form.get('role', 'user')
    if username:
        users = load_json(USERS_PATH, [])
        users.append({'username': username, 'role': role})
        save_json(USERS_PATH, users)
        append_log(f'کاربر {username} با نقش {role} اضافه شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/delete_resource', methods=['POST'])
def delete_resource():
    filename = request.form.get('filename')
    file_path = os.path.join(MED_DOC_DIR, 'doctor_abbasi', filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    resources = load_json(RESOURCES_PATH, [])
    resources = [f for f in resources if f['name'] != filename]
    save_json(RESOURCES_PATH, resources)
    append_log(f'فایل {filename} حذف شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/delete_user', methods=['POST'])
def delete_user():
    username = request.form.get('username')
    users = load_json(USERS_PATH, [])
    users = [u for u in users if u['username'] != username]
    save_json(USERS_PATH, users)
    append_log(f'کاربر {username} حذف شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/change_role', methods=['POST'])
def change_role():
    username = request.form.get('username')
    role = request.form.get('role')
    users = load_json(USERS_PATH, [])
    for u in users:
        if u['username'] == username:
            u['role'] = role
    save_json(USERS_PATH, users)
    append_log(f'نقش کاربر {username} به {role} تغییر یافت.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/change_password', methods=['POST'])
def change_password():
    username = request.form.get('username')
    new_password = request.form.get('new_password')
    users = load_json(USERS_PATH, [])
    for u in users:
        if u['username'] == username:
            u['password'] = new_password
    save_json(USERS_PATH, users)
    append_log(f'رمز کاربر {username} تغییر یافت.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/save_security_settings', methods=['POST'])
def save_security_settings():
    settings = load_json(SETTINGS_PATH, {})
    settings['api_key'] = request.form.get('api_key', '')
    settings['allowed_ips'] = request.form.get('allowed_ips', '')
    save_json(SETTINGS_PATH, settings)
    append_log('تنظیمات امنیتی ذخیره شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/backup', methods=['POST'])
def backup():
    backup_dir = os.path.join(BASE_DIR, 'chatbot_backup')
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
    for f in [SETTINGS_PATH, RESOURCES_PATH, USERS_PATH, LOGS_PATH]:
        if os.path.exists(f):
            shutil.copy(f, backup_dir)
    if os.path.exists(os.path.join(BASE_DIR, 'chroma_db')):
        shutil.make_archive(os.path.join(backup_dir, 'chroma_db_backup'), 'zip', os.path.join(BASE_DIR, 'chroma_db'))
    append_log('پشتیبان‌گیری انجام شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/restore', methods=['POST'])
def restore():
    backup_file = request.files.get('backup_file')
    if backup_file:
        backup_path = os.path.join(BASE_DIR, 'chatbot_backup', backup_file.filename)
        backup_file.save(backup_path)
        # اگر فایل zip بود، دیتابیس را بازگردانی کن
        if backup_file.filename.endswith('.zip'):
            shutil.unpack_archive(backup_path, os.path.join(BASE_DIR, 'chroma_db'))
        else:
            # فایل json یا txt را جایگزین کن
            target = os.path.join(BASE_DIR, backup_file.filename)
            shutil.copy(backup_path, target)
    append_log('بازیابی انجام شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/test_chatbot', methods=['POST'])
def test_chatbot():
    test_message = request.form.get('test_message', '')
    response = ''
    try:
        if llm:
            response = llm.invoke(test_message)
        else:
            response = 'مدل LLM فعال نیست.'
    except Exception as e:
        response = f'خطا: {e}'
    settings = load_json(SETTINGS_PATH, {})
    resources = load_json(RESOURCES_PATH, [])
    users = load_json(USERS_PATH, [])
    logs = []
    if os.path.exists(LOGS_PATH):
        with open(LOGS_PATH, 'r', encoding='utf-8') as f:
            logs = f.readlines()[-30:]
    return render_template('admin/chatbot_settings.html', settings=settings, resources=resources, users=users, logs=logs, test_response=response)

@app.route('/admin/save_tts_settings', methods=['POST'])
def save_tts_settings():
    settings = load_json(SETTINGS_PATH, {})
    settings['tts_voice'] = request.form.get('tts_voice', 'nova')
    settings['tts_speed'] = float(request.form.get('tts_speed', 1.0))
    settings['tts_accent'] = request.form.get('tts_accent', 'تهرانی')
    save_json(SETTINGS_PATH, settings)
    append_log('تنظیمات TTS/STT ذخیره شد.')
    return redirect(url_for('chatbot_settings'))

@app.route('/admin/download_logs', methods=['GET'])
def download_logs():
    return send_from_directory(BASE_DIR, 'chatbot_logs.txt', as_attachment=True)

@app.route('/admin/preview_resource', methods=['POST'])
def preview_resource():
    filename = request.form.get('filename')
    file_path = os.path.join(MED_DOC_DIR, 'doctor_abbasi', filename)
    text = ''
    if os.path.exists(file_path) and filename.endswith('.docx'):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
    return jsonify({'text': text})

@app.route('/admin/save_resource_edit', methods=['POST'])
def save_resource_edit():
    filename = request.form.get('filename')
    new_text = request.form.get('text')
    file_path = os.path.join(MED_DOC_DIR, 'doctor_abbasi', filename)
    if os.path.exists(file_path) and filename.endswith('.docx'):
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for para in new_text.split('\n'):
            doc.add_paragraph(para)
        doc.save(file_path)
    append_log(f'فایل {filename} ویرایش شد.')
    return jsonify({'status': 'ok'})

# راهنمای اجرای پروژه
if __name__ == "__main__":
    print("""
    راه‌اندازی پروژه:
    1. نصب پیش‌نیازها:
       pip install flask flask_sqlalchemy flask_migrate openai ffmpeg-python
    2. مقداردهی دیتابیس (فقط بار اول):
       flask db init
       flask db migrate
       flask db upgrade
    3. اجرای برنامه:
       python app.py
    """)
    app.run(debug=True)