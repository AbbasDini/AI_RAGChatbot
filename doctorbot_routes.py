from flask import Blueprint, render_template, request, redirect, url_for, jsonify, send_file, flash, session
from doctorbot_models import db, DoctorBotSettings, MedicalDocument
import os
import docx
import requests
from werkzeug.utils import secure_filename
import uuid
from flask import send_from_directory
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from doctors_data import doctors_info

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.vectorstores import Chroma
from llm_utils import get_llm




UPLOAD_FOLDER = 'Med_doc/'
ALLOWED_EXTENSIONS = {'doc', 'docx'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

doctorbot_bp = Blueprint('doctorbot', __name__, url_prefix='/doctorbot')

@doctorbot_bp.route('/settings', methods=['GET', 'POST'])
def doctorbot_settings():
    doctor_list = [(k, v['name']) for k, v in doctors_info.items()]
    if request.method == 'POST':
        llm_model = request.form.get('llm_model')
        tts_model = request.form.get('tts_model')
        stt_model = request.form.get('stt_model')
        embedding_model = request.form.get('embedding_model')
        selected_doctor = request.form.get('doctor_name')
        # حذف مقداردهی api_key از env و عدم ذخیره آن در دیتابیس
        settings = DoctorBotSettings.query.first()
        if not settings:
            settings = DoctorBotSettings()
        settings.llm_model = llm_model
        settings.tts_model = tts_model
        settings.stt_model = stt_model
        settings.embedding_model = embedding_model
        # settings.api_key = حذف شود
        db.session.add(settings)
        db.session.commit()
        # آپلود و embedding فایل Word
        file = request.files.get('doc_file')
        if file and allowed_file(file.filename) and selected_doctor:
            filename = secure_filename(file.filename or "uploaded.docx")
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            # استخراج متن از Word
            doc = docx.Document(filepath)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            # ارسال به مدل embedding AvalAI
            from llm_utils import AVALAI_API_KEY
            embedding = get_avalai_embedding(full_text, embedding_model, AVALAI_API_KEY)
            if embedding is None:
                flash('خطا در دریافت embedding از سرویس. سند ذخیره نشد.', 'danger')
                return redirect(url_for('doctorbot.doctorbot_settings'))
            # جلوگیری از درج رکورد تکراری
            existing_doc = MedicalDocument.query.filter_by(filename=filename, doctor_name=selected_doctor).first()
            if existing_doc:
                flash('این سند قبلاً برای این پزشک ثبت شده است.', 'warning')
                return redirect(url_for('doctorbot.doctorbot_settings'))
            # ذخیره در دیتابیس اصلی
            med_doc = MedicalDocument()
            med_doc.filename = filename
            med_doc.content = full_text
            med_doc.embedding = embedding
            med_doc.doctor_name = selected_doctor
            db.session.add(med_doc)
            db.session.commit()
            # --- درج در دیتابیس Chroma ---
            try:
                from langchain_core.documents import Document
                BASE_DIR = os.path.dirname(os.path.abspath(__file__))
                CHROMA_DB_DIR = os.path.join(BASE_DIR, "chroma_db")
                doctor_attrs = {"doctor": selected_doctor, "source": filepath}
                doc_obj = Document(page_content=full_text, metadata=doctor_attrs)
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                chunks = text_splitter.split_documents([doc_obj])
                embedding_model_local = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
                db_chroma = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embedding_model_local)
                db_chroma.add_documents(chunks)
                db_chroma.persist()
                flash('سند به دیتابیس Chroma اضافه شد.', 'success')
            except Exception as e:
                flash(f'خطا در افزودن به ChromaDB: {e}', 'danger')
            # --- پایان درج در دیتابیس Chroma ---
            flash('سند پزشکی با موفقیت افزوده شد.', 'success')
        else:
            if file:
                flash('فرمت فایل یا پزشک انتخابی مجاز نیست.', 'danger')
        flash('تنظیمات با موفقیت ذخیره شد.', 'success')
        return redirect(url_for('doctorbot.doctorbot_settings'))
    return render_template('doctorbot_settings.html', doctor_list=doctor_list)


def get_avalai_embedding(text, embedding_model, api_key):
    # فرض بر این است که AvalAI یک endpoint embedding دارد
    url = 'https://api.avalai.ir/v1/embeddings'
    headers = {'Authorization': f'Bearer {api_key}'}
    payload = {
        'model': embedding_model,
        'input': text
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        if response.status_code == 200:
            return response.json().get('embedding')
    except Exception as e:
        print('خطا در دریافت embedding:', e)
    return None

@doctorbot_bp.route('/select_doctor', methods=['GET', 'POST'])
def doctorbot_select_doctor():
    doctor_list = [(k, v['name']) for k, v in doctors_info.items()]
    if request.method == 'POST':
        selected_doctor = request.form.get('doctor_name')
        if selected_doctor:
            session['selected_doctor'] = selected_doctor
            return redirect(url_for('doctorbot.doctorbot_chat'))
    return render_template('doctorbot_select_doctor.html', doctor_list=doctor_list)

@doctorbot_bp.route('/chat', methods=['GET'])
def doctorbot_chat():
    if 'selected_doctor' not in session:
        return redirect(url_for('doctorbot.doctorbot_select_doctor'))
    doctor_id = session['selected_doctor']
    doctor = doctors_info.get(doctor_id, {})
    settings = DoctorBotSettings.query.first()
    # دریافت اسناد مربوط به پزشک منتخب
    med_docs = MedicalDocument.query.filter_by(doctor_name=doctor_id).all()
    return render_template('doctorbot_chat.html', doctor=doctor, settings=settings, med_docs=med_docs)

@doctorbot_bp.route('/api/chat', methods=['POST'])
def doctorbot_api_chat():
    data = request.get_json()
    user_message = data.get('message', '')
    # دریافت تنظیمات فعال
    settings = DoctorBotSettings.query.first()
    if not settings:
        return jsonify({'response': 'تنظیمات چت‌بات یافت نشد.'})
    # دریافت پزشک انتخابی
    selected_doctor = session.get('selected_doctor')
    if not selected_doctor:
        return jsonify({'response': 'پزشک انتخاب نشده است.'})
    # دریافت embedding اسناد پزشکی فقط برای پزشک منتخب
    med_docs = MedicalDocument.query.filter_by(doctor_name=selected_doctor).all()
    context = '\n'.join([doc.content for doc in med_docs]) if med_docs else ''
    # ساخت prompt ترکیبی
    prompt = f"""
شما یک دستیار پزشکی حرفه‌ای هستید. با توجه به اطلاعات زیر از اسناد پزشکی و سوال کاربر، به صورت خلاصه و دقیق و با زبان فارسی و قالب Markdown پاسخ دهید.

اطلاعات اسناد:
{context}

سوال کاربر:
{user_message}
"""
    # فراخوانی مدل LLM داینامیک
    llm = get_llm()
    llm_response = None
    if llm:
        try:
            llm_response = llm.invoke(prompt)
        except Exception as e:
            print('خطا در فراخوانی LLM:', e)
    return jsonify({'response': llm_response or 'پاسخی از مدل دریافت نشد.'})


def call_avalai_llm(prompt, llm_model, api_key):
    # تعیین endpoint و پارامترها بر اساس مدل انتخابی
    if llm_model.startswith('avalai-'):
        url = 'https://api.avalai.ir/v1/chat/completions'
        headers = {'Authorization': f'Bearer {api_key}'}
        payload = {
            'model': llm_model,
            'messages': [
                {'role': 'system', 'content': 'شما یک دستیار پزشکی حرفه‌ای هستید.'},
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 512,
            'temperature': 0.2
        }
    elif llm_model.startswith('gpt-'):
        url = 'https://api.openai.com/v1/chat/completions'
        headers = {'Authorization': f'Bearer ' + api_key, 'Content-Type': 'application/json'}
        payload = {
            'model': llm_model,
            'messages': [
                {'role': 'system', 'content': 'You are a professional medical assistant.'},
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 512,
            'temperature': 0.2
        }
    elif llm_model.startswith('meta-llama') or llm_model.startswith('nousresearch/'):
        # فرض بر این که API سازگار با OpenRouter یا HuggingFace است (نیاز به پیاده‌سازی واقعی دارد)
        url = 'https://openrouter.ai/api/v1/chat/completions'
        headers = {'Authorization': f'Bearer {api_key}', 'HTTP-Referer': 'https://your-app.com', 'X-Title': 'DoctorBot'}
        payload = {
            'model': llm_model,
            'messages': [
                {'role': 'system', 'content': 'You are a professional medical assistant.'},
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 512,
            'temperature': 0.2
        }
    else:
        return None
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=60)
        print('LLM raw response:', response.status_code, response.text)
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and result['choices']:
                return result['choices'][0]['message']['content']
    except Exception as e:
        print('خطا در فراخوانی LLM:', e)
    return None

@doctorbot_bp.route('/api/tts', methods=['POST'])
def doctorbot_api_tts():
    data = request.get_json()
    text = data.get('text', '')
    settings = DoctorBotSettings.query.first()
    if not settings or not text:
        return jsonify({'error': 'تنظیمات یا متن یافت نشد.'}), 400
    audio_path = avalai_tts(text, settings.tts_model, settings.api_key)
    if audio_path:
        audio_url = url_for('doctorbot.doctorbot_tts_file', filename=os.path.basename(audio_path))
        return jsonify({'audio_url': audio_url})
    return jsonify({'error': 'خطا در تولید صوت.'}), 500

@doctorbot_bp.route('/tts_audio/<filename>')
def doctorbot_tts_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

def avalai_tts(text, tts_model, api_key):
    url = 'https://api.avalai.ir/v1/audio/tts'
    headers = {'Authorization': f'Bearer {api_key}'}
    payload = {
        'model': tts_model,
        'input': text,
        'voice': 'female',
        'response_format': 'wav'
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=60)
        if response.status_code == 200:
            filename = f"tts_{uuid.uuid4().hex}.wav"
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            with open(filepath, 'wb') as f:
                f.write(response.content)
            return filepath
    except Exception as e:
        print('خطا در TTS AvalAI:', e)
    return None

@doctorbot_bp.route('/api/stt', methods=['POST'])
def doctorbot_api_stt():
    if 'audio' not in request.files:
        return jsonify({'error': 'فایل صوتی ارسال نشده است.'}), 400
    audio_file = request.files['audio']
    settings = DoctorBotSettings.query.first()
    if not settings:
        return jsonify({'error': 'تنظیمات یافت نشد.'}), 400
    text = avalai_stt(audio_file, settings.stt_model, settings.api_key)
    if text:
        return jsonify({'text': text})
    return jsonify({'error': 'خطا در تبدیل صوت به متن.'}), 500

def avalai_stt(audio_file, stt_model, api_key):
    url = 'https://api.avalai.ir/v1/audio/stt'
    headers = {'Authorization': f'Bearer {api_key}'}
    files = {'file': (audio_file.filename, audio_file.stream, audio_file.mimetype)}
    data = {'model': stt_model, 'language': 'fa'}
    try:
        response = requests.post(url, data=data, files=files, headers=headers, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result.get('text')
    except Exception as e:
        print('خطا در STT AvalAI:', e)
    return None

@doctorbot_bp.route('/api/upload_doc', methods=['POST'])
def doctorbot_api_upload_doc():
    # آپلود و embedding فایل Word
    return jsonify({'status': 'success'}) 