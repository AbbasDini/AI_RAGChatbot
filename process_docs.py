import os
import re
# hazm import removed
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.vectorstores import Chroma
import pdfplumber
from docx import Document as DocxDocument

# Define paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MED_DOC_DIR = os.path.join(BASE_DIR, "Med_doc")
CHROMA_DB_DIR = os.path.join(BASE_DIR, "chroma_db") # Directory to store ChromaDB

# Define doctor attributes for metadata
# These attributes should match the folder names in Med_doc
DOCTOR_ATTRIBUTES = {
    "doctor_abbasi": {
        "city": "hamedan",
        "specialty": "neurologist",
        "experience": "10+"
    },
    "doctor_mohammadi": {
        "city": "hamedan",
        "specialty": "neurologist",
        "experience": "10+"
    },
    "doctor_sharifi": {
        "city": "shiraz",
        "specialty": "oncologist",
        "experience": "10+"
    }
    # Add other doctors here with their attributes
}
print('=== SCRIPT STARTED ===')
import sys
print(sys.version)

print('--- Start process_docs.py ---')
try:
    import os
    print('os imported')
    import re
    print('re imported')
    # hazm import removed
    try:
        from pdf2image import convert_from_path
        print('pdf2image imported')
        import pytesseract
        print('pytesseract imported')
        OCR_AVAILABLE = True
    except ImportError:
        print('OCR libraries not available')
        OCR_AVAILABLE = False
    from langchain.document_loaders import PyPDFLoader
    print('PyPDFLoader imported')
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    print('RecursiveCharacterTextSplitter imported')
    from langchain.embeddings import SentenceTransformerEmbeddings
    print('SentenceTransformerEmbeddings imported')
    from langchain.vectorstores import Chroma
    print('Chroma imported')
except Exception as e:
    print(f'Import error: {e}')


def clean_persian_text(text):
    # حذف کاراکترهای کنترلی و نویزهای رایج OCR و غیرمجاز
    text = re.sub(r'[\u200c\u200f\u202a-\u202e]', '', text)
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)  # حذف کاراکترهای غیرقابل چاپ
    text = re.sub(r'[■]', '', text)  # حذف نویزهای OCR
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[a-zA-Z0-9]', '', text)  # حذف حروف و اعداد لاتین
    return text.strip()

def extract_text_with_pdfplumber(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = "\n".join(page.extract_text() or '' for page in pdf.pages)
        return text
    except Exception as e:
        print(f"pdfplumber extraction failed for {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path):
    try:
        doc = DocxDocument(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"docx extraction failed for {docx_path}: {e}")
        return ""

# Function to load and process documents for a specific doctor
def process_doctor_docs(doctor_name: str):
    print(f'Processing doctor: {doctor_name}')
    doctor_folder = os.path.join(MED_DOC_DIR, doctor_name)
    if not os.path.exists(doctor_folder):
        print(f"Error: Folder for {doctor_name} not found at {doctor_folder}")
        return []
        
    # Get attributes for this doctor
    doctor_attrs = DOCTOR_ATTRIBUTES.get(doctor_name, {}) # Get attributes, default to empty dict
    if not doctor_attrs:
        print(f"Warning: Attributes not found for doctor {doctor_name}. No additional metadata will be added.")

    documents = []
    for filename in os.listdir(doctor_folder):
        filepath = os.path.join(doctor_folder, filename)
        if filename.endswith(".docx"):
            print(f"Loading {filepath} (Word)")
            text = extract_text_from_docx(filepath)
            if text and len(text.strip()) > 50:
                text = clean_persian_text(text)
                from langchain_core.documents import Document
                doc = Document(page_content=text, metadata={"doctor": doctor_name, "source": filepath, **doctor_attrs})
                docs = [doc]
                print(f'Loaded 1 doc from {filepath} using python-docx')
            else:
                docs = []
        else:
            docs = []
        documents.extend(docs)

    print(f"Loaded {len(documents)} pages from {doctor_name}'s documents.")
    return documents

# Main processing logic
if __name__ == "__main__":
    print('--- Main processing logic started ---')
    all_documents = []
    # Get list of doctors (subdirectories in Med_doc) that are also in DOCTOR_ATTRIBUTES
    doctor_folders_to_process = [d for d in os.listdir(MED_DOC_DIR) 
                                 if os.path.isdir(os.path.join(MED_DOC_DIR, d))
                                 and d in DOCTOR_ATTRIBUTES] # Only process doctors with defined attributes
    print(f"Found doctors with defined attributes to process: {doctor_folders_to_process}")
    
    # Warn about doctor folders without defined attributes
    all_doctor_folders = [d for d in os.listdir(MED_DOC_DIR) 
                          if os.path.isdir(os.path.join(MED_DOC_DIR, d))]
    doctors_without_attrs = [d for d in all_doctor_folders if d not in DOCTOR_ATTRIBUTES]
    if doctors_without_attrs:
        print(f"Warning: The following doctor folders exist but do not have defined attributes in DOCTOR_ATTRIBUTES and will NOT be processed: {doctors_without_attrs}")

    for doctor in doctor_folders_to_process:
        print(f'--- Processing {doctor} ---')
        all_documents.extend(process_doctor_docs(doctor))
        print(f'--- Finished {doctor} ---')

    print(f'Total loaded documents: {len(all_documents)}')
    if not all_documents:
        print("No documents found to process based on defined attributes. Please check the Med_doc folder and DOCTOR_ATTRIBUTES.")
    else:
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=150)
        print(f"Splitting {len(all_documents)} documents into chunks...")
        chunks = text_splitter.split_documents(all_documents)
        print(f"Created {len(chunks)} chunks.")

        # ذخیره اولین chunk استخراج‌شده از erfani.pdf برای بازبینی
        for chunk in chunks:
            if 'doctor_abbasi' in chunk.metadata.get('doctor', '') and 'erfani.pdf' in chunk.metadata.get('source', ''):
                with open('sample_chunk_erfani.txt', 'w', encoding='utf-8') as f:
                    f.write(chunk.page_content)
                print('Sample chunk from erfani.pdf saved to sample_chunk_erfani.txt')
                break

        # Create embeddings using a better multilingual model for Persian
        print("Creating embeddings...")
        embedding_model = SentenceTransformerEmbeddings(model_name="paraphrase-multilingual-MiniLM-L12-v2")
        print("Embedding model loaded.")

        # Initialize and populate ChromaDB (This will overwrite existing data if directory exists)
        print(f"Initializing and populating ChromaDB at {CHROMA_DB_DIR}...")
        # Ensure the directory exists or ChromaDB handles creation/overwriting
        # It's safer to delete the old directory before running this script if you make changes to metadata structure
        
        # Optional: Delete existing ChromaDB directory to ensure clean rebuild
        # import shutil
        # if os.path.exists(CHROMA_DB_DIR):
        #     print(f"Deleting existing ChromaDB directory: {CHROMA_DB_DIR}")
        #     shutil.rmtree(CHROMA_DB_DIR)
        #     print("Deleted existing ChromaDB directory.")

        # Create the directory if it doesn't exist
        if not os.path.exists(CHROMA_DB_DIR):
            os.makedirs(CHROMA_DB_DIR)

        db = Chroma.from_documents(chunks, embedding_model, persist_directory=CHROMA_DB_DIR)
        print("ChromaDB populated successfully with metadata.")

        print("Document processing complete.")  